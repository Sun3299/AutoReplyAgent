"""
文档解析器 - 支持多种文档格式

文档解析是 RAG 系统的第一步，负责将各种格式的文档转换为文本内容。

支持格式：
- TXT: 纯文本文件（.txt, .text, .md）
- PDF: PDF 文档（通过 PyPDF2）
- DOCX: Word 文档（通过 python-docx）
- HTML: 网页文件（通过 BeautifulSoup）

设计模式：
- 工厂模式：DocumentLoaderFactory 统一管理各种加载器
- 策略模式：每种文档格式有独立的加载策略
- 适配器模式：将不同格式转换为统一的 (content, source) 格式

依赖安装：
    pip install PyPDF2 python-docx beautifulsoup4
"""

import re
from abc import ABC, abstractmethod
from typing import List, Optional
from pathlib import Path


# 尝试导入可选依赖
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None


class DocumentLoader(ABC):
    """
    文档加载器抽象基类
    
    定义文档加载器的接口，所有具体加载器都需要实现：
    - load(): 加载文档
    - is_supported(): 检查是否支持某文件类型
    """
    
    @abstractmethod
    def load(self, file_path: str) -> List[tuple]:
        """
        加载文档
        
        Returns:
            List[(content, source)]
            - content: 提取的文本内容
            - source: 来源标识（如文件名）
        """
        pass
    
    @abstractmethod
    def is_supported(self, file_path: str) -> bool:
        """
        检查是否支持此文件类型
        
        Args:
            file_path: 文件路径
            
        Returns:
            bool: 如果支持返回 True
        """
        pass


class TextLoader(DocumentLoader):
    """
    TXT 文本加载器
    
    加载纯文本文件，支持 .txt, .text, .md 等格式。
    简单处理：直接读取文件内容，清理多余空白字符。
    """
    
    def is_supported(self, file_path: str) -> bool:
        """检查是否为文本文件"""
        return Path(file_path).suffix.lower() in ['.txt', '.text', '.md']
    
    def load(self, file_path: str) -> List[tuple]:
        """
        加载文本文件
        
        Returns:
            List[(content, source)]
        """
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # 清理文本
        content = self._clean_text(content)
        
        # 返回 (内容, 文件名)
        return [(content, Path(file_path).name)]
    
    def _clean_text(self, text: str) -> str:
        """
        清理文本
        
        处理：
        1. 统一换行符（\r\n, \r -> \n）
        2. 合并多余空行（3+ 换行 -> 2 换行）
        3. 合并多余空格/制表符
        4. 去除首尾空白
        """
        # 统一换行符
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)
        
        # 合并多余空行
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # 合并空格
        text = re.sub(r'[ \t]+', ' ', text)
        
        return text.strip()


class PDFLoader(DocumentLoader):
    """
    PDF 文档加载器
    
    按页提取 PDF 内容，每页作为一个文本片段。
    """
    
    def is_supported(self, file_path: str) -> bool:
        """检查是否为 PDF 文件"""
        return Path(file_path).suffix.lower() == '.pdf'
    
    def load(self, file_path: str) -> List[tuple]:
        """
        加载 PDF 文档
        
        按页提取文本，每页返回一个元组。
        
        Returns:
            List[(page_content, source)]
            - page_content: 单页内容
            - source: "文件名#页码"
        """
        if PdfReader is None:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
        
        results = []
        reader = PdfReader(file_path)
        
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text = self._clean_text(text)
                if text.strip():
                    # 格式：文件名#页码
                    source = f"{Path(file_path).name}#p{page_num+1}"
                    results.append((text, source))
        
        return results
    
    def _clean_text(self, text: str) -> str:
        """
        清理 PDF 文本
        
        PDF 文本提取经常有格式问题，需要特殊处理。
        """
        # 在大小写字母数字之间加空格（如 "ABC123" -> "ABC 123"）
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        
        # 在数字和字母之间加空格（如 "1.A" -> "1. A"）
        text = re.sub(r'(\d+)\.([A-Z])', r'\1. \2', text)
        
        # 合并多余空行
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # 合并空格
        text = re.sub(r'[ \t]+', ' ', text)
        
        return text.strip()


class DocxLoader(DocumentLoader):
    """
    Word 文档加载器
    
    加载 .docx 文件，按段落和表格提取内容。
    标题段落单独返回，正文段落合并返回。
    """
    
    def is_supported(self, file_path: str) -> bool:
        """检查是否为 Word 文件"""
        return Path(file_path).suffix.lower() in ['.docx', '.doc']
    
    def load(self, file_path: str) -> List[tuple]:
        """
        加载 Word 文档
        
        处理流程：
        1. 遍历所有段落
        2. 标题段落（Heading）单独作为一个片段
        3. 非标题段落合并为一个片段
        4. 提取所有表格内容
        
        Returns:
            List[(content, source)]
        """
        if DocxDocument is None:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        results = []
        doc = DocxDocument(file_path)
        file_name = Path(file_path).name
        
        # 用于合并的段落
        current_paragraphs = []
        
        # 遍历段落
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                # 空行，保存当前合并的段落
                if current_paragraphs:
                    content = ' '.join(current_paragraphs)
                    if content.strip():
                        results.append((content, file_name))
                    current_paragraphs = []
                continue
            
            # 检查是否为标题
            if para.style.name.startswith('Heading'):
                # 是标题，先保存当前合并的段落
                if current_paragraphs:
                    content = ' '.join(current_paragraphs)
                    if content.strip():
                        results.append((content, file_name))
                    current_paragraphs = []
                
                # 标题单独保存
                results.append((text, f"{file_name}#{para.style.name}"))
            else:
                # 非标题，添加到合并队列
                current_paragraphs.append(text)
        
        # 保存最后一段
        if current_paragraphs:
            content = ' '.join(current_paragraphs)
            if content.strip():
                results.append((content, file_name))
        
        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    results.append((' | '.join(cells), f"{file_name}#table"))
        
        return results
    
    def _clean_text(self, text: str) -> str:
        """清理文本"""
        return text.strip()


class HtmlLoader(DocumentLoader):
    """
    HTML 文档加载器
    
    加载网页文件，提取标题、段落、列表等内容。
    """
    
    def is_supported(self, file_path: str) -> bool:
        """检查是否为 HTML 文件"""
        return Path(file_path).suffix.lower() in ['.html', '.htm']
    
    def load(self, file_path: str) -> List[tuple]:
        """
        加载 HTML 文档
        
        提取内容：
        1. 标题 (h1-h6)
        2. 段落 (p)
        3. 列表项 (li)
        
        注意：脚本和样式内容会被移除。
        
        Returns:
            List[(content, source)]
        """
        if BeautifulSoup is None:
            raise ImportError("beautifulsoup4 not installed. Run: pip install beautifulsoup4")
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        
        # 移除脚本和样式
        for script in soup(['script', 'style']):
            script.decompose()
        
        results = []
        file_name = Path(file_path).name
        
        # 提取标题
        for header in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            text = header.get_text().strip()
            if text:
                results.append((text, f"{file_name}#{header.name}"))
        
        # 提取段落（长度大于10）
        for para in soup.find_all('p'):
            text = para.get_text().strip()
            if text and len(text) > 10:
                results.append((text, file_name))
        
        # 提取列表项（长度大于5）
        for li in soup.find_all('li'):
            text = li.get_text().strip()
            if text and len(text) > 5:
                results.append((text, f"{file_name}#li"))
        
        return results


class DocumentLoaderFactory:
    """
    文档加载器工厂
    
    统一管理各种文档加载器，根据文件类型自动选择合适的加载器。
    
    使用示例：
        factory = DocumentLoaderFactory()
        docs = factory.load("path/to/document.pdf")
        
        # 或使用便捷函数
        docs = load_document("path/to/document.pdf")
    """
    
    def __init__(self):
        """初始化工厂，注册所有加载器"""
        self.loaders: List[DocumentLoader] = [
            TextLoader(),
            PDFLoader(),
            DocxLoader(),
            HtmlLoader(),
        ]
    
    def load(self, file_path: str) -> List[tuple]:
        """
        加载文档
        
        自动选择合适的加载器。
        
        Args:
            file_path: 文档路径
            
        Returns:
            List[(content, source)]
            
        Raises:
            ValueError: 不支持的文件类型
        """
        for loader in self.loaders:
            if loader.is_supported(file_path):
                return loader.load(file_path)
        
        raise ValueError(f"Unsupported file type: {file_path}")
    
    def load_directory(
        self,
        directory: str,
        extensions: Optional[List[str]] = None,
        recursive: bool = True
    ) -> List[tuple]:
        """
        加载目录下所有指定类型的文档
        
        Args:
            directory: 目录路径
            extensions: 要加载的文件扩展名列表
            recursive: 是否递归子目录
            
        Returns:
            所有文档的 (content, source) 列表
        """
        if extensions is None:
            extensions = ['.txt', '.pdf', '.docx', '.html', '.htm', '.md']
        
        results = []
        path = Path(directory)
        
        # glob 模式
        pattern = "**/*" if recursive else "*"
        
        for file_path in path.glob(pattern):
            # 只处理文件，且扩展名匹配
            if file_path.is_file() and file_path.suffix.lower() in extensions:
                try:
                    docs = self.load(str(file_path))
                    results.extend(docs)
                except Exception as e:
                    print(f"Error loading {file_path}: {e}")
        
        return results


def load_document(file_path: str) -> List[tuple]:
    """
    便捷函数：加载单个文档
    
    Args:
        file_path: 文档路径
        
    Returns:
        List[(content, source)]
    """
    factory = DocumentLoaderFactory()
    return factory.load(file_path)


def load_documents(directory: str, recursive: bool = True) -> List[tuple]:
    """
    便捷函数：加载目录下所有文档
    
    Args:
        directory: 目录路径
        recursive: 是否递归子目录
        
    Returns:
        所有文档的 (content, source) 列表
    """
    factory = DocumentLoaderFactory()
    return factory.load_directory(directory, recursive=recursive)
